#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Insere logos SES | CIEVS-MT | Rede CIEVS no cabeçalho do DOCX institucional.

Entrada padrão: docs/SECRETARIA_ESTADO_SAUDE_CIEVS_MT.docx
Saída:          docs/CIEVS_MT_com_logos.docx
"""
from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
LOGO_DIR = ROOT / "assets" / "logos"
DEFAULT_IN = ROOT / "docs" / "SECRETARIA_ESTADO_SAUDE_CIEVS_MT.docx"
DEFAULT_OUT = ROOT / "docs" / "CIEVS_MT_com_logos.docx"

# PNG preferidos (python-docx embute bem); SVG fica de fallback textual se ausente.
LOGOS = [
    ("SES-MT", LOGO_DIR / "ses_governo_mt_banner.png", LOGO_DIR / "ses_mt_oficial.png"),
    ("CIEVS-MT", LOGO_DIR / "cievs_ses_faixa_preta.png", LOGO_DIR / "cievs_mt.svg"),
    ("Rede CIEVS", LOGO_DIR / "cievs_rede_ses_banner.png", None),
]

LOGO_WIDTH_CM = 3.6
TITLE_LINE = "SES-MT · CIEVS-MT · Radar LACEN"
FOOTER_LINE = (
    "Documento institucional — sem microdado nominal · "
    "Circulação interna CIEVS/VE"
)


def _resolve_logo(primary: Path, fallback: Path | None) -> Path | None:
    if primary.exists() and primary.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif"}:
        return primary
    if fallback and fallback.exists() and fallback.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif"}:
        return fallback
    if primary.exists() and primary.suffix.lower() == ".svg":
        # SVG não embute de forma confiável no Word via python-docx
        return None
    return None


def _set_run_font(run, *, size_pt: float = 9, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "Calibri"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    if color is not None:
        run.font.color.rgb = color


def _clear_paragraphs(container) -> None:
    for p in list(container.paragraphs):
        p._element.getparent().remove(p._element)


def _add_header(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    _clear_paragraphs(header)

    # Tabela 1×3 para alinhar as três logomarcas
    table = header.add_table(rows=1, cols=3, width=Cm(16.5))
    table.autofit = True
    for idx, (caption, primary, fallback) in enumerate(LOGOS):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        path = _resolve_logo(primary, fallback)
        if path is not None:
            run = p.add_run()
            run.add_picture(str(path), width=Cm(LOGO_WIDTH_CM))
        else:
            run = p.add_run(caption)
            _set_run_font(run, size_pt=10, bold=True, color=RGBColor(0x1B, 0x32, 0x81))
        # Legenda discreta
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = p2.add_run(caption)
        _set_run_font(cap, size_pt=7, color=RGBColor(0x5A, 0x6A, 0x85))

    title = header.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(2)
    tr = title.add_run(TITLE_LINE)
    _set_run_font(tr, size_pt=10, bold=True, color=RGBColor(0x1B, 0x32, 0x81))


def _add_footer(doc: Document, se_ref: str | None = None) -> None:
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    _clear_paragraphs(footer)

    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    se_bit = f" · {se_ref}" if se_ref else ""
    run = p.add_run(f"{FOOTER_LINE} · {date.today().isoformat()}{se_bit}")
    _set_run_font(run, size_pt=7, color=RGBColor(0x5A, 0x6A, 0x85))


def _detect_se(doc: Document) -> str | None:
    for para in doc.paragraphs[:12]:
        t = para.text or ""
        if "SE " in t.upper() or "Semana Epidemiológica" in t or "SE30" in t.replace(" ", ""):
            # tenta extrair padrão SE 30/2026 ou 2026-SE30
            import re

            m = re.search(r"SE\s*(\d{1,2})\s*/\s*(\d{4})", t, re.I)
            if m:
                return f"SE {m.group(1)}/{m.group(2)}"
            m = re.search(r"(\d{4})-SE(\d{1,2})", t, re.I)
            if m:
                return f"SE {m.group(2)}/{m.group(1)}"
    return None


def inserir_logos(src: Path, dst: Path) -> Path:
    if not src.exists():
        raise FileNotFoundError(f"DOCX de entrada não encontrado: {src}")
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(src))
    se_ref = _detect_se(doc)
    _add_header(doc)
    _add_footer(doc, se_ref=se_ref)
    doc.save(str(dst))
    return dst


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--in", dest="src", type=Path, default=DEFAULT_IN)
    ap.add_argument("--out", dest="dst", type=Path, default=DEFAULT_OUT)
    args = ap.parse_args()
    out = inserir_logos(args.src, args.dst)
    print(f"OK: {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
